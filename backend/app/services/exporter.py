from docx import Document
import markdown
import pdfkit

class MinutesExporter:
    def export_as_docx(self, minutes, output_path):
        doc = Document()
        doc.add_heading('Classroom Minutes', 0)
        
        # Parse and add sections
        sections = minutes.split('\n\n')
        for section in sections:
            if section.startswith('Main Topics'):
                doc.add_heading('Main Topics', 1)
                doc.add_paragraph(section.replace('Main Topics:', ''))
            # ... etc
        
        doc.save(output_path)
        return output_path
    
    def export_as_html(self, minutes, output_path):
        html = markdown.markdown(minutes)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"""
            <html><body>
            <h1>Classroom Minutes</h1>
            {html}
            </body></html>
            """)
        return output_path