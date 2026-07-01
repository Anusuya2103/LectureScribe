# backend/app/services/export_service.py
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import json
from datetime import datetime
import re
import os
from pathlib import Path
from deep_translator import GoogleTranslator

class ExportService:
    """Export transcripts with multiple formats and translation."""
    
    def __init__(self):
        try:
            self.translator = GoogleTranslator()
        except:
            self.translator = None
    
    def export_pdf(self, transcript: str, segments: List[Dict], metadata: Dict, filename: str) -> str:
        """Export to PDF with timestamps."""
        pdf = FPDF()
        pdf.add_page()
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Classroom Minutes - Transcript", ln=1, align='C')
        pdf.ln(5)
        
        # Metadata
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(200, 6, txt=f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1, align='C')
        pdf.cell(200, 6, txt=f"Duration: {metadata.get('duration', 0):.2f} seconds", ln=1, align='C')
        pdf.ln(5)
        
        # Content with timestamps
        pdf.set_font("Arial", size=10)
        for seg in segments:
            timestamp = f"[{self._format_time(seg.get('start', 0))}]"
            pdf.set_font("Arial", 'B', 9)
            pdf.cell(35, 6, txt=timestamp, ln=0)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(155, 6, txt=seg.get('text', ''))
            pdf.ln(1)
        
        os.makedirs("exports", exist_ok=True)
        output_path = f"exports/{filename}.pdf"
        pdf.output(output_path)
        return output_path
    
    def export_pdf_without_timestamps(self, transcript: str, metadata: Dict, filename: str) -> str:
        """Export PDF without timestamps."""
        pdf = FPDF()
        pdf.add_page()
        
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Classroom Minutes - Transcript", ln=1, align='C')
        pdf.ln(5)
        
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(200, 6, txt=f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1, align='C')
        pdf.ln(5)
        
        clean_text = self._remove_timestamps(transcript)
        pdf.set_font("Arial", size=10)
        pdf.multi_cell(190, 6, txt=clean_text)
        
        os.makedirs("exports", exist_ok=True)
        output_path = f"exports/{filename}_clean.pdf"
        pdf.output(output_path)
        return output_path
    
    def export_docx(self, transcript: str, segments: List[Dict], metadata: Dict, filename: str) -> str:
        """Export to DOCX with timestamps."""
        doc = Document()
        
        title = doc.add_heading('Classroom Minutes - Transcript', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Duration: {metadata.get('duration', 0):.2f} seconds")
        doc.add_paragraph()
        
        for seg in segments:
            p = doc.add_paragraph()
            timestamp = f"[{self._format_time(seg.get('start', 0))}]"
            run = p.add_run(timestamp)
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(f"  {seg.get('text', '').strip()}")
            run.font.size = Pt(10)
        
        os.makedirs("exports", exist_ok=True)
        output_path = f"exports/{filename}.docx"
        doc.save(output_path)
        return output_path
    
    def export_bilingual_pdf(self, tamil_text: str, english_text: str, metadata: Dict, filename: str) -> str:
        """Export bilingual PDF with Tamil and English side-by-side."""
        pdf = FPDF()
        pdf.add_page()
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Classroom Minutes - Bilingual Transcript", ln=1, align='C')
        pdf.ln(5)
        
        # Metadata
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(200, 6, txt=f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1, align='C')
        pdf.cell(200, 6, txt=f"Duration: {metadata.get('duration', 0):.2f} seconds", ln=1, align='C')
        pdf.ln(5)
        
        # Split into sentences for side-by-side
        tamil_sentences = tamil_text.split('. ')
        english_sentences = english_text.split('. ')
        
        pdf.set_font("Arial", size=10)
        for i in range(max(len(tamil_sentences), len(english_sentences))):
            tamil_part = tamil_sentences[i] + '.' if i < len(tamil_sentences) else ''
            english_part = english_sentences[i] + '.' if i < len(english_sentences) else ''
            
            if tamil_part:
                pdf.set_font("Arial", 'B', 10)
                pdf.cell(95, 6, txt="[Tamil]", ln=0)
                pdf.set_font("Arial", size=10)
                pdf.multi_cell(95, 6, txt=tamil_part)
            
            if english_part:
                pdf.set_font("Arial", 'B', 10)
                pdf.cell(95, 6, txt="[English]", ln=0)
                pdf.set_font("Arial", size=10)
                pdf.multi_cell(95, 6, txt=english_part)
            
            pdf.ln(2)
        
        os.makedirs("exports", exist_ok=True)
        output_path = f"exports/{filename}_bilingual.pdf"
        pdf.output(output_path)
        return output_path
    
    def export_translated(self, transcript: str, segments: List[Dict], 
                          metadata: Dict, filename: str, target_lang: str = "en") -> Dict[str, str]:
        """Export with translation."""
        results = {}
        
        translated_text = self._translate_text(transcript, target_lang)
        
        # Create DOCX with side-by-side
        doc = Document()
        title = doc.add_heading('Classroom Minutes - Bilingual Transcript', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Table for side-by-side
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Original'
        hdr_cells[1].text = f'Translation ({target_lang.upper()})'
        
        for seg in segments:
            row_cells = table.add_row().cells
            row_cells[0].text = seg.get('text', '')
            translated_seg = self._translate_text(seg.get('text', ''), target_lang)
            row_cells[1].text = translated_seg
        
        os.makedirs("exports", exist_ok=True)
        docx_path = f"exports/{filename}_translated.docx"
        doc.save(docx_path)
        results['docx'] = docx_path
        
        txt_path = f"exports/{filename}_translated.txt"
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("=== ORIGINAL ===\n\n")
            f.write(transcript)
            f.write("\n\n=== TRANSLATION ===\n\n")
            f.write(translated_text)
        results['txt'] = txt_path
        
        return results
    
    def export_summary(self, summary: str, structured_minutes: Dict, filename: str) -> str:
        """Export summary as DOCX."""
        doc = Document()
        
        title = doc.add_heading('Classroom Minutes - Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        doc.add_heading('📝 Summary', level=1)
        doc.add_paragraph(summary)
        doc.add_paragraph()
        
        doc.add_heading('📑 Structured Minutes', level=1)
        for key, value in structured_minutes.items():
            if value:
                doc.add_heading(f'▸ {key.replace("_", " ").title()}', level=2)
                if isinstance(value, list):
                    for item in value:
                        doc.add_paragraph(f"• {item}", style='List Bullet')
                else:
                    doc.add_paragraph(str(value))
                doc.add_paragraph()
        
        os.makedirs("exports", exist_ok=True)
        output_path = f"exports/{filename}_summary.docx"
        doc.save(output_path)
        return output_path
    
    def _translate_text(self, text: str, target_lang: str = "en") -> str:
        """Translate text using deep-translator."""
        if not text or not text.strip():
            return text
        
        try:
            if self.translator:
                translated = self.translator.translate(text, dest=target_lang)
                return translated
            return text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
    
    def _format_time(self, seconds: float) -> str:
        """Format seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def _remove_timestamps(self, text: str) -> str:
        """Remove timestamps from text."""
        return re.sub(r'\[\d+:\d+\.\d+\]\s*', '', text)